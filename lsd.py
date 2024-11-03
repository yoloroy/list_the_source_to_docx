#!/usr/bin/env python3
import re

import docx
from docx.text.run import Run


def list_the_source_to_docx(
        destination: str,
        source_files: list[str],
        source_style: Run,
        title_heading_level: int
):
    destination_doc = docx.Document()

    for source_file in source_files:
        destination_doc.add_heading(source_file, title_heading_level)
        print("adding", source_file)
        with open(source_file, 'r') as f:
            paragraph = destination_doc.add_paragraph()
            _add_run_copy(paragraph, source_style, f.read())
        destination_doc.add_page_break()
        print(source_file, "added")

    destination_doc.save(destination)


def _add_run_copy(paragraph, run, text=None):
    # source: https://github.com/python-openxml/python-docx/issues/519

    r = paragraph.add_run(text=run.text if text is None else text, style=run.style)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.color.theme_color = run.font.color.theme_color
    #r.font.color.type = run.font.color.type
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    return r


def _source_files_by_regex(regex: re.Pattern[str]) -> list[str]:
    current_working_directory = os.getcwd()

    return [
        file[len(current_working_directory) + 1:]
        for root, dirs, files in os.walk(current_working_directory)
        for file in (f'{root}/{file}' for file in files)
        if regex.match(file) and '/.' not in file
    ]


if __name__ == '__main__':
    from argparse import ArgumentParser
    from sys import argv
    import os
    import re

    if argv[0].endswith('.py'):
        argv = argv[1:]

    dir_path = os.path.dirname(os.path.realpath(__file__))

    parser = ArgumentParser(
        prog="List the Source to Docs : lsd",
        description="This program creates single .docx file from the source code provided"
    )
    parser.add_argument(
        '-d', '--destination',
        default='./listing.docx',
        help='.docx file where result shall be put'
    )
    parser.add_argument(
        '-f', '--source_files',
        default=_source_files_by_regex(re.compile('(.*/)*[a-z_]+\\.py')),
        type=lambda arg: _source_files_by_regex(re.compile(arg)),
        help='Regexp filter of files that shall be included into listing'
    )
    parser.add_argument(
        '-t', '--title_heading_level',
        default=1,
        type=int,
        help="Heading level used for titling styling"
    )
    parser.add_argument(
        '-s', '--source_style',
        default=docx.Document(dir_path + "/_source_style.docx").paragraphs[0].runs[0],
        type=lambda path: docx.Document(path).paragraphs[0].runs[0],
        help=".docx file, which first symbol will be used as source code style reference"
    )

    args = parser.parse_args(argv)
    if argv != ['-h'] and argv != ['--help']:
        list_the_source_to_docx(**vars(args))
